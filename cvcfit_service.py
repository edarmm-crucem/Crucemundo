import io
import re
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from config import BARCOS_MAP, DRIVE_ROOT_ID
from services_google import (
    find_child_folder,
    find_file_by_name,
    get_sheet_titles,
    get_single_cell,
    get_range,
)

def normalize_text(value):
    if value is None:
        return ""
    return str(value).strip().lower()

def extract_first_number(value):
    if value is None:
        return 0
    m = re.search(r"(\d+)", str(value))
    return int(m.group(1)) if m else 0

def safe_filename(text):
    text = re.sub(r'[\\/:*?"<>|]', "", str(text))
    text = re.sub(r"\s+", " ", text).strip()
    return text

def first_line(value):
    if value is None:
        return ""
    return str(value).splitlines()[0].strip()

def parse_nombre_apellidos_from_g24(g24_value):
    raw = first_line(g24_value)
    if "/" in raw:
        nombre, apellidos = raw.split("/", 1)
        return nombre.strip(), apellidos.strip()
    return raw.strip(), ""

def parse_locator(locator):
    locator = normalize_text(locator).upper().replace(" ", "")
    m = re.fullmatch(r"([A-Z]+)(\d{2})(\d{2})(\d{2})-(\d{3})", locator)
    if not m:
        raise Exception("El localizador debe tener formato BARCOAAMMDD-XXX, por ejemplo ALB260101-001.")
    code, yy, mm, dd, seq = m.groups()
    if code not in BARCOS_MAP:
        raise Exception(f"Código de barco no reconocido: {code}")
    full_boat = BARCOS_MAP[code]
    departure_name = f"{full_boat}_{yy}{mm}{dd}"
    year_4 = f"20{yy}"
    fecha_salida = datetime.strptime(f"{year_4}-{mm}-{dd}", "%Y-%m-%d").date()
    return {
        "locator": locator,
        "boat_code": code,
        "boat_name": full_boat,
        "year_folder": year_4,
        "departure_file_name": departure_name,
        "fecha_salida": fecha_salida,
        "fecha_limite_pago": fecha_salida - timedelta(days=30),
    }

def find_drive_file_for_locator(locator_info):
    year_folder = find_child_folder(DRIVE_ROOT_ID, locator_info["year_folder"])
    if not year_folder:
        raise Exception(f"No existe la carpeta del año {locator_info['year_folder']}.")
    boat_folder = find_child_folder(year_folder["id"], locator_info["boat_name"])
    if not boat_folder:
        raise Exception(f"No existe la carpeta del barco {locator_info['boat_name']} en {locator_info['year_folder']}.")
    target_file = find_file_by_name(boat_folder["id"], locator_info["departure_file_name"])
    if not target_file:
        raise Exception(f"No existe el archivo {locator_info['departure_file_name']}.")
    return {
        "id": target_file["id"],
        "name": target_file["name"],
        "url": target_file.get("webViewLink") or f"https://docs.google.com/spreadsheets/d/{target_file['id']}/edit",
    }

def build_money_text(matrix):
    lines = []
    for row in matrix:
        cleaned = [str(c).strip() for c in row if str(c).strip()]
        if cleaned:
            lines.append(" | ".join(cleaned))
    return "\n".join(lines).strip()

def underline_paragraph(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)

def build_cvc_fit_from_locator(locator):
    locator_info = parse_locator(locator)
    drive_file = find_drive_file_for_locator(locator_info)
    spreadsheet_id = drive_file["id"]

    sheet_titles = get_sheet_titles(spreadsheet_id)
    if locator_info["locator"] not in sheet_titles:
        raise Exception(f"No existe una pestaña con nombre {locator_info['locator']} dentro del spreadsheet.")

    sheet_title = locator_info["locator"]
    g24 = get_single_cell(spreadsheet_id, sheet_title, "G24")
    p24 = get_single_cell(spreadsheet_id, sheet_title, "P24")

    nombre, apellidos = parse_nombre_apellidos_from_g24(g24)
    dni = first_line(p24)

    values_people = [
        get_single_cell(spreadsheet_id, sheet_title, "G22"),
        get_single_cell(spreadsheet_id, sheet_title, "K22"),
        get_single_cell(spreadsheet_id, sheet_title, "N22"),
        get_single_cell(spreadsheet_id, sheet_title, "P22"),
    ]
    personas = sum(extract_first_number(v) for v in values_people)

    values_rooms = [
        get_single_cell(spreadsheet_id, sheet_title, "G20"),
        get_single_cell(spreadsheet_id, sheet_title, "K20"),
        get_single_cell(spreadsheet_id, sheet_title, "N20"),
        get_single_cell(spreadsheet_id, sheet_title, "P20"),
    ]
    habitaciones = sum(extract_first_number(v) for v in values_rooms)

    dinero = get_range(spreadsheet_id, sheet_title, "G33:R53")
    total = get_single_cell(spreadsheet_id, sheet_title, "Q55")
    fecha_salida = locator_info["fecha_salida"]
    fecha_limite_pago = locator_info["fecha_limite_pago"]

    filename = safe_filename(
        f"CVC Fit {apellidos} {nombre} {locator_info['boat_name']} salida {fecha_salida.strftime('%d %m')}.docx"
    )

    return {
        "locator": locator_info["locator"],
        "boat_name": locator_info["boat_name"],
        "spreadsheet_id": spreadsheet_id,
        "spreadsheet_url": drive_file["url"],
        "sheet_title": sheet_title,
        "nombre": nombre,
        "apellidos": apellidos,
        "dni": dni,
        "personas": personas,
        "habitaciones": habitaciones,
        "dinero_matrix": dinero,
        "dinero_text": build_money_text(dinero),
        "total": total,
        "fecha_salida": fecha_salida,
        "fecha_salida_str": fecha_salida.strftime("%d/%m/%Y"),
        "fecha_limite_pago": fecha_limite_pago,
        "fecha_limite_pago_str": fecha_limite_pago.strftime("%d/%m/%Y"),
        "filename": filename,
    }

def build_cvc_fit_doc(data):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    def add_line(text="", bold=False, align=None, size=None):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        r = p.add_run(text)
        r.bold = bold
        if size:
            r.font.size = Pt(size)
        return p

    def add_section_title(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(11)
        return p

    def add_blank_line():
        doc.add_paragraph()

    add_line("Lugar y fecha: ________________________________", align=WD_ALIGN_PARAGRAPH.RIGHT)
    title = add_line("CONTRATO DE VIAJE COMBINADO", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    underline_paragraph(title)
    add_blank_line()

    add_section_title("DATOS DEL VIAJERO")
    add_line(f"Nombre: {data['nombre']}")
    add_line(f"Apellidos: {data['apellidos']}")
    add_line(f"DNI / Pasaporte: {data['dni']}")
    add_line(f"Nº Personas: {data['personas']}")
    add_line(f"Nº Habitaciones: {data['habitaciones']}")
    add_blank_line()

    add_section_title("DATOS DEL VIAJE")
    add_line(f"Barco: {data['boat_name']}")
    add_line(f"Fecha de salida: {data['fecha_salida_str']}")
    add_line(f"Fecha límite pago: {data['fecha_limite_pago_str']}")
    add_line(f"Total: {data['total']}")
    add_blank_line()

    add_section_title("PRECIO Y FORMA DE PAGO")
    if data["dinero_text"]:
        for linea in data["dinero_text"].splitlines():
            add_line(linea)
    else:
        add_line("(Sin detalle extraído del rango G33:R53)")

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
